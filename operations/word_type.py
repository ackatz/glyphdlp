import base64
from docx import Document
from io import BytesIO
from app.detections import regex
import re


def scan(input) -> dict:
    results = {}
    doc = Document(BytesIO(input))

    # Scan paragraphs for plaintext emails and hyperlinks
    for para in doc.paragraphs:
        for name, pattern in regex.patterns.items():
            matches = re.findall(pattern, para.text)
            if matches:
                if name not in results:
                    results[name] = []
                results[name].extend(matches)

    # Scan tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for name, pattern in regex.patterns.items():
                    matches = re.findall(pattern, cell.text)
                    if matches:
                        if name not in results:
                            results[name] = []
                        results[name].extend(matches)

    # Check for hyperlinks (including mailto addresses)
    for rel in doc.part.rels.values():
        if "mailto:" in rel.target_ref:
            for name, pattern in regex.patterns.items():
                matches = re.findall(pattern, rel.target_ref)
                if matches:
                    if name not in results:
                        results[name] = []
                    results[name].extend(matches)

    return results


def redact(input) -> str:
    doc = Document(BytesIO(input))

    # Redact paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            for name, pattern in regex.patterns.items():
                run.text = re.sub(pattern, "[REDACTED]", run.text)

    # Redact tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for name, pattern in regex.patterns.items():
                            run.text = re.sub(pattern, "[REDACTED]", run.text)

    # Write the modified document to bytes
    output_stream = BytesIO()
    doc.save(output_stream)
    redacted_doc_bytes = output_stream.getvalue()

    # Convert the bytes to Base64-encoded string
    redacted_doc_base64 = base64.b64encode(redacted_doc_bytes).decode("utf-8")

    return redacted_doc_base64
